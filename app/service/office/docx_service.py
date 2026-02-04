

import logging
import os
from cgitb import text

import gradio as gr
from docx import Document
from docx.shared import RGBColor
from langchain_text_splitters import CharacterTextSplitter

from app.service.langchain.index import LLM_obj
from app.service.langchain.openai_llm_langchain_service import \
    OpenaiLLMLangchainClass
from app.service.langchain.siliconflow_llm_langchain_service import \
    SiliconflowLLMLangchainClass
from app.util.config import Settings


async def extract_text_with_properties(doc):
    """提取文档中所有的文本及其属性"""
    text_properties = []
    for para in doc.paragraphs:
        for run in para.runs:
            text_properties.append({
                'text': run.text,
                'font': run.font.name,
                'size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None,
                'style': para.style.name
            })
    return text_properties


async def restore_text_to_docx(source_properties, process_text):
    """恢复文档原有的格式"""
    new_doc = Document()
    current_index = 0
    for prop in source_properties:
        end_index = current_index + len(prop['text'])
        next_text = process_text[current_index:end_index]

        paragraph = new_doc.add_paragraph()
        run = paragraph.add_run(next_text)

        # 应用原始格式
        # run.font.name = prop['font']    #字体可能不通用
        run.font.size = prop['size']
        run.font.bold = prop['bold']
        run.font.italic = prop['italic']
        run.font.underline = prop['underline']
        if prop['color']:
            run.font.color.rgb = RGBColor.from_string(prop['color'])
        paragraph.style = prop['style']
        current_index = end_index

    return new_doc


async def save_docx(new_doc: Document, out_path: str):
    """保存文档"""
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    new_doc.save(out_path)


async def launch_translation_docx(file_path: str, prompt: str, struct: str, output_name: str, language_radio: str, llm_radio: str, process=gr.Progress(track_tqdm=True)) -> str:
    try:
        llm_obj: OpenaiLLMLangchainClass | SiliconflowLLMLangchainClass = LLM_obj(
            llm_radio, 'str')
        out_path = Settings.FILE_OUTPUT_PATH+output_name

        # 读取文档
        doc = Document(file_path)
        text_properties = await extract_text_with_properties(doc)

        # 合并文本
        full_text = "".join([prop['text'] for prop in text_properties])
        logging.info(f"docx+++++++++++翻译前内容:{full_text}")

        # # 分割文本
        # text_spliter = CharacterTextSplitter(
        #     chunk_size=Settings.CHUNK_LENGTH, separator="\n")
        # chunks = text_spliter.split_text(full_text)

        # # llm 翻译
        # process_chunks = [await (llm_obj.translates_chain(prompt)).
        #                   chain_ainvoke(chunk, struct, language_radio)
        #                   for chunk in chunks]
        # process_text = "".join(process_chunks)

        process_text = await (llm_obj.translates_chain(prompt)).chain_ainvoke(full_text, struct, language_radio)
        # 恢复文档原有的格式
        # new_doc = await restore_text_to_docx(text_properties,process_text)

        new_doc = Document()
        logging.info(f"docx==================翻译后内容:{process_text}")

        new_doc.add_paragraph(process_text.content)
        new_doc.save(out_path)

        # 保存文档
        # await save_docx(new_doc, out_path)
    except Exception as e:
        logging.error("------翻译word文档出错------:"+str(e))
    return out_path
